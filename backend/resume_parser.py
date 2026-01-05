# backend/resume_parser.py
import re
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image
import io

# --- robust text extractor for PDF or DOCX file path ---
def extract_text_from_file(path_or_file):
    """
    Accepts either a path string (saved file) or a file-like object (Streamlit upload).
    Returns extracted text (string).
    """
    text = ""
    try:
        # If it's a path string (common in our pages), open accordingly
        if isinstance(path_or_file, str):
            if path_or_file.lower().endswith(".pdf"):
                reader = PdfReader(path_or_file)
                for page in reader.pages:
                    ptxt = page.extract_text()
                    if ptxt:
                        text += ptxt + "\n"
            elif path_or_file.lower().endswith(".docx"):
                doc = Document(path_or_file)
                for para in doc.paragraphs:
                    if para.text:
                        text += para.text + "\n"
        else:
            # file-like (Streamlit upload)
            if hasattr(path_or_file, "type") and "pdf" in path_or_file.type:
                reader = PdfReader(path_or_file)
                for page in reader.pages:
                    ptxt = page.extract_text()
                    if ptxt:
                        text += ptxt + "\n"
            else:
                # try docx
                try:
                    doc = Document(path_or_file)
                    for para in doc.paragraphs:
                        if para.text:
                            text += para.text + "\n"
                except Exception:
                    pass
    except Exception:
        pass
    return (text or "").strip()

# short alias for readability
def extract_text_from_pdf(path_or_file):
    return extract_text_from_file(path_or_file)

# --- extract email, name, links ---
def extract_user_details(text):
    text = text or ""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    # email
    email = None
    m = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    if m:
        email = m.group(0)

    # links
    links = re.findall(r"(https?://[^\s,;]+)", text)

    # name detection: search first 6 lines for a likely name
    name = None
    for ln in lines[:8]:
        ln_clean = re.sub(r"[^A-Za-z\s]", "", ln).strip()
        low = ln.lower()
        # skip headings
        if any(skip in low for skip in ["resume", "curriculum", "objective", "profile", "skills", "contact"]):
            continue
        words = ln_clean.split()
        if 1 < len(words) <= 4 and all(w.isalpha() for w in words):
            # plausible name
            name = " ".join([w.capitalize() for w in words])
            break

    return {"name": name or None, "email": email or None, "links": links or []}

# --- skills detection using a curated list ---
SKILLS_DB = [
    "python","sql","excel","power bi","tableau","pandas","numpy","machine learning",
    "deep learning","tensorflow","pytorch","data analysis","statistics","r","matlab",
    "html","css","javascript","react","node.js","java","c++","c#","django","flask","aws","azure","git"
]

def extract_skills_from_text(text):
    t = (text or "").lower()
    found = []
    for s in SKILLS_DB:
        if re.search(rf"\b{re.escape(s)}\b", t):
            found.append(s)
    # dedupe
    out = []
    for x in found:
        if x not in out:
            out.append(x)
    return out

# wrapper naming consistency
def extract_skills(text):
    return extract_skills_from_text(text)

# --- optional: extract first image from PDF (returns PIL Image) ---
def extract_image_from_pdf_safe(path_or_file):
    try:
        if isinstance(path_or_file, str) and path_or_file.lower().endswith(".pdf"):
            reader = PdfReader(path_or_file)
        else:
            reader = PdfReader(path_or_file)
        for page in reader.pages:
            if "/XObject" in page["/Resources"]:
                xObject = page["/Resources"]["/XObject"].get_object()
                for obj in xObject:
                    try:
                        if xObject[obj]["/Subtype"] == "/Image":
                            data = xObject[obj].get_data()
                            img = Image.open(io.BytesIO(data))
                            return img
                    except Exception:
                        continue
    except Exception:
        return None
    return None
