# backend/resume_template.py
from docx import Document
from docx.shared import Pt
import io

def generate_docx_template_bytes(name, email, links, job_title, skills_list, summary_text):
    doc = Document()
    # Title (name)
    h = doc.add_heading(level=0)
    r = h.add_run(name or "Candidate Name")
    r.bold = True
    r.font.size = Pt(20)

    # Contact
    if email:
        doc.add_paragraph(f"Email: {email}")
    if links:
        doc.add_paragraph("Links:")
        for l in links:
            doc.add_paragraph(f"- {l}")

    # Summary
    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(summary_text or f"Applying for {job_title}")

    # Skills
    doc.add_heading("Skills", level=2)
    for s in skills_list:
        doc.add_paragraph(f"- {s}")

    # Experience placeholder
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("• Your role — Company — Dates\n• Achievements and responsibilities (quantify where possible).")

    # Education
    doc.add_heading("Education", level=2)
    doc.add_paragraph("• Degree — Institution — Year")

    # Projects
    doc.add_heading("Projects", level=2)
    doc.add_paragraph("• Project 1 — short description and tech used")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()
