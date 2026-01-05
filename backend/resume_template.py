from docx import Document
from io import BytesIO

def generate_docx_template_bytes(name, email, links, job_title, skills, summary):
    doc = Document()

    doc.add_heading(f"Resume - {name}", 0)

    doc.add_heading("Contact Info", level=1)
    doc.add_paragraph(f"Email: {email}")
    
    if links:
        doc.add_paragraph("Links:")
        for link in links:
            doc.add_paragraph(f"- {link}")

    doc.add_heading("Job Title", level=1)
    doc.add_paragraph(job_title)

    doc.add_heading("Recommended Skills", level=1)
    for skill in skills:
        doc.add_paragraph(f"- {skill}")

    doc.add_heading("Summary / Feedback", level=1)
    doc.add_paragraph(summary)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
