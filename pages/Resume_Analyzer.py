import streamlit as st
import os
import json
from datetime import datetime
import PyPDF2
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import re
from io import BytesIO

# ========= PREDEFINED ROLES, JDs & SKILLS ========= #

JOB_PROFILES = {
    "Software Engineer": {
        "jd": """We are looking for a Software Engineer responsible for designing, developing, 
and maintaining high-quality applications. You will work with cross-functional teams, 
write clean and scalable code, perform code reviews, and contribute to the full SDLC.""",
        "skills": [
            "Python", "Java", "C++", "Data Structures", "Algorithms", "Object Oriented Programming",
            "Git", "REST APIs", "SQL", "Unit Testing", "Agile", "Debugging"
        ],
    },
    "Full Stack Developer": {
        "jd": """As a Full Stack Developer you will build end-to-end web applications, 
work on both frontend and backend, integrate APIs, and ensure good performance and security.""",
        "skills": [
            "HTML", "CSS", "JavaScript", "React", "Node.js", "Express", "REST APIs",
            "MongoDB", "SQL", "Git", "Responsive Design", "Authentication"
        ],
    },
    "Frontend Developer": {
        "jd": """We need a Frontend Developer to create responsive, user-friendly interfaces, 
optimize web pages for speed, and collaborate with designers and backend engineers.""",
        "skills": [
            "HTML", "CSS", "JavaScript", "React", "Redux",
            "Responsive Design", "Cross-Browser Compatibility", "Figma", "UI Development"
        ],
    },
    "Backend Developer": {
        "jd": """Backend Developer will design scalable APIs, manage databases, 
implement business logic, and ensure performance and security on the server side.""",
        "skills": [
            "Python", "Django", "Flask", "Node.js", "REST APIs", "SQL",
            "PostgreSQL", "MySQL", "Database Design", "Authentication", "Docker"
        ],
    },
    "Data Scientist": {
        "jd": """Data Scientist will build models, analyze large datasets, 
generate insights, and help business stakeholders make data-driven decisions.""",
        "skills": [
            "Python", "R", "Statistics", "Machine Learning", "Pandas", "NumPy",
            "Scikit-learn", "Data Visualization", "SQL", "Feature Engineering"
        ],
    },
    "Data Analyst": {
        "jd": """Data Analyst will clean and analyze data, build dashboards, 
prepare reports, and support decision-making with data insights.""",
        "skills": [
            "Excel", "SQL", "Power BI", "Tableau", "Data Cleaning", "Data Visualization",
            "Reporting", "Pivot Tables", "Basic Statistics"
        ],
    },
    "Machine Learning Engineer": {
        "jd": """ML Engineer will design, build, and deploy machine learning models into production, 
optimize performance, and collaborate with data scientists and engineers.""",
        "skills": [
            "Python", "Scikit-learn", "TensorFlow", "PyTorch", "Machine Learning",
            "Model Deployment", "MLOps", "Docker", "APIs", "Data Pipelines"
        ],
    },
    "DevOps Engineer": {
        "jd": """DevOps Engineer will manage CI/CD pipelines, automate deployments, 
monitor systems, and ensure reliability and scalability of infrastructure.""",
        "skills": [
            "Linux", "Bash", "CI/CD", "Jenkins", "Docker", "Kubernetes",
            "AWS", "Azure", "Monitoring", "Git", "Terraform"
        ],
    },
    "Cloud Engineer": {
        "jd": """Cloud Engineer will design, deploy, and manage cloud infrastructure, 
ensure security and cost optimization, and support development teams.""",
        "skills": [
            "AWS", "Azure", "GCP", "Virtual Machines", "VPC", "Cloud Security",
            "IAM", "Docker", "Kubernetes", "Networking", "Monitoring"
        ],
    },
    "Product Manager": {
        "jd": """Product Manager will own product roadmap, gather requirements, 
work with cross-functional teams, and ensure successful product delivery and adoption.""",
        "skills": [
            "Product Roadmap", "User Stories", "Stakeholder Management",
            "Market Research", "Wireframing", "Analytics", "Agile", "Prioritization"
        ],
    },
    "Project Manager": {
        "jd": """Project Manager will plan, execute, and close projects, 
manage timelines, resources, risks, and communicate with stakeholders.""",
        "skills": [
            "Project Planning", "Scheduling", "Risk Management", "Stakeholder Management",
            "MS Project", "JIRA", "Agile", "Scrum", "Communication"
        ],
    },
    "Business Analyst": {
        "jd": """Business Analyst will gather requirements, map processes, 
analyze business problems, and propose data-driven solutions.""",
        "skills": [
            "Requirements Gathering", "Process Mapping", "SQL", "Documentation",
            "Stakeholder Communication", "UML", "User Stories", "Gap Analysis"
        ],
    },
    "Sales Executive": {
        "jd": """Sales Executive will identify leads, pitch products, 
follow up with clients, and close deals to achieve revenue targets.""",
        "skills": [
            "Lead Generation", "Cold Calling", "Negotiation", "CRM", "Customer Relationship",
            "Sales Pitch", "Objection Handling", "Closing Deals", "Communication"
        ],
    },
    "Sales Manager": {
        "jd": """Sales Manager will manage sales team, define targets, monitor performance, 
and drive strategies to increase revenue and market share.""",
        "skills": [
            "Sales Strategy", "Team Management", "Pipeline Management", "CRM",
            "Forecasting", "Negotiation", "Target Setting", "Coaching", "Reporting"
        ],
    },
    "Inside Sales Representative": {
        "jd": """Inside Sales Representative will handle inbound and outbound calls, 
qualify leads, nurture prospects, and schedule demos/meetings.""",
        "skills": [
            "CRM", "Cold Calling", "Lead Qualification", "Email Outreach",
            "Communication", "Objection Handling", "Follow-ups"
        ],
    },
    "Digital Marketing Specialist": {
        "jd": """Digital Marketing Specialist will plan and execute online campaigns, 
optimize SEO/SEM, manage social media, and track performance metrics.""",
        "skills": [
            "SEO", "SEM", "Google Ads", "Facebook Ads", "Content Marketing",
            "Email Marketing", "Google Analytics", "Social Media Management"
        ],
    },
    "HR Manager": {
        "jd": """HR Manager will handle recruitment, employee engagement, performance management, 
and ensure HR policies and compliance.""",
        "skills": [
            "Recruitment", "Interviewing", "Onboarding", "Performance Management",
            "Employee Engagement", "HR Policies", "Conflict Resolution"
        ],
    },
    "UI/UX Designer": {
        "jd": """UI/UX Designer will create user-centered designs, wireframes, prototypes, 
and collaborate with engineers to implement intuitive interfaces.""",
        "skills": [
            "Figma", "Wireframing", "Prototyping", "User Research",
            "Usability Testing", "UI Design", "Design Systems"
        ],
    },
    "QA Engineer": {
        "jd": """QA Engineer will design and execute test plans, write test cases, 
and ensure product quality through manual and automated testing.""",
        "skills": [
            "Test Cases", "Test Planning", "Manual Testing", "Automation Testing",
            "Selenium", "Bug Tracking", "JIRA", "Regression Testing"
        ],
    },
    "Customer Support Specialist": {
        "jd": """Customer Support Specialist will resolve customer queries, troubleshoot issues, 
and ensure high customer satisfaction through timely support.""",
        "skills": [
            "Customer Support", "Ticketing Systems", "Communication",
            "Problem Solving", "Email Support", "Chat Support", "Phone Support"
        ],
    },
    "Financial Analyst": {
        "jd": """Financial Analyst will analyze financial data, create reports, 
build models, and support budgeting and forecasting.""",
        "skills": [
            "Financial Modeling", "Excel", "Forecasting", "Budgeting",
            "Reporting", "Power BI", "Variance Analysis"
        ],
    },
}

# Union of all skills (to detect skills present in resume)
ALL_SKILLS = sorted({skill for v in JOB_PROFILES.values() for skill in v["skills"]})


# ========= HELPERS ========= #

def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip().lower()


def extract_text_and_image(file_bytes: bytes, filename: str):
    """
    Read resume bytes and return (text, image_bytes or None).
    Photo extraction is implemented for DOCX. For PDFs, only text for now.
    """
    text = ""
    image_bytes = None
    fname = filename.lower()

    if fname.endswith(".pdf"):
        reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception:
                continue
        image_bytes = None

    elif fname.endswith(".docx"):
        bio = BytesIO(file_bytes)
        doc = Document(bio)
        text = "\n".join([p.text for p in doc.paragraphs])

        # Try to extract first embedded image if any
        try:
            for rel in doc.part.rels.values():
                if rel.reltype == RT.IMAGE:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    break
        except Exception:
            image_bytes = None

    elif fname.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="ignore")
        image_bytes = None

    else:
        text = ""
        image_bytes = None

    return text, image_bytes


def extract_skills_from_text(text, skills_universe):
    """Return list of skills from skills_universe that appear in text (case-insensitive)."""
    text_lower = normalize(text)
    matched = []
    for skill in skills_universe:
        pattern = r"\b" + re.escape(skill.lower()) + r"\b"
        if re.search(pattern, text_lower):
            matched.append(skill)
    return matched


def guess_name_from_email(email: str) -> str:
    """Guess a name from the email local-part (before @)."""
    local = email.split("@")[0]
    local = re.sub(r"[_\.]+", " ", local)  # replace _ and . with space
    parts = [p for p in local.split() if p.isalpha()]
    if not parts:
        return ""
    return " ".join(p.capitalize() for p in parts)


def extract_basic_details(text: str, candidate_name_input: str = ""):
    """
    Extract details: name, email, phone, links from resume text.
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    email = None
    phone = None
    links = []

    # Email
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    if email_match:
        email = email_match.group(0)

    # Phone (simple, allows +91, spaces, dashes, brackets)
    phone_match = re.search(r"(\+?\d[\d \-\(\)]{8,}\d)", text)
    if phone_match:
        phone = phone_match.group(1)

    # Links
    links = re.findall(r"(https?://\S+)", text)

    # Name heuristic priority:
    # 1. user input
    # 2. from top lines
    # 3. from email local part
    name = candidate_name_input.strip() if candidate_name_input else ""

    if not name:
        for line in lines[:7]:  # look at first few lines
            lower = line.lower()
            if any(x in lower for x in ["resume", "curriculum vitae", "cv", "@", "http", "www."]):
                continue
            if sum(c.isalpha() for c in line) < 3:
                continue
            words = line.split()
            if 1 <= len(words) <= 4:
                if sum(w[0].isupper() for w in words if w and w[0].isalpha()) >= 1:
                    name = line
                    break

    if not name and email:
        name = guess_name_from_email(email)

    return {
        "name": name,
        "email": email or "",
        "phone": phone or "",
        "links": links,
    }


def append_result_to_json(file_path, record):
    """Append a single record dict to JSON list file."""
    if os.path.exists(file_path):
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                if not isinstance(data, list):
                    data = []
        except Exception:
            data = []
    else:
        data = []

    data.append(record)
    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def categorize_role(job_title: str) -> str:
    """Roughly map job title to a template category."""
    t = job_title.lower()

    if any(x in t for x in ["sales", "business development", "account manager"]):
        return "sales"
    if any(x in t for x in ["data scientist", "data analyst", "ml", "machine learning", "ai"]):
        return "data"
    if any(x in t for x in ["developer", "engineer", "devops", "cloud", "software", "frontend", "backend", "full stack"]):
        return "tech"
    if any(x in t for x in ["product manager", "project manager", "scrum", "business analyst"]):
        return "pm_ba"
    if any(x in t for x in ["ui/ux", "ux", "designer"]):
        return "design"
    if any(x in t for x in ["support", "customer success", "helpdesk"]):
        return "support"
    if any(x in t for x in ["finance", "financial", "accountant"]):
        return "finance"
    if any(x in t for x in ["hr", "human resources", "talent acquisition"]):
        return "hr"
    return "generic"


def generate_resume_template(job_title, candidate_name, matched_skills, missing_skills, jd_text: str = ""):
    """
    Create a tailored resume template text for the selected job.
    - Template changes based on job_title category.
    - JD text is used to guide the summary section.
    """
    name_display = candidate_name if candidate_name else "Your Name"
    all_skills_for_section = matched_skills + missing_skills
    skills_line = ", ".join(all_skills_for_section) if all_skills_for_section else f"{job_title} core skills"

    # Use first meaningful line of JD as a one-line role context if available
    jd_summary_line = ""
    if jd_text:
        lines = [l.strip() for l in jd_text.splitlines() if l.strip()]
        if lines:
            sentence = lines[0]
            if len(sentence) > 10:
                jd_summary_line = sentence

    category = categorize_role(job_title)

    # ==== SUMMARY & EXPERIENCE BLOCKS BY CATEGORY ====
    if category == "tech":
        summary = f"""
SUMMARY
{jd_summary_line or f"{job_title} with strong fundamentals in software engineering and problem solving."}
Experienced in building, testing, and maintaining reliable applications and services.
Comfortable working with modern development practices, version control, and agile teams.
"""
        experience_block = f"""
PROFESSIONAL EXPERIENCE
Company Name | {job_title} | Location | MM/YYYY – Present
• Design, develop, and maintain applications following clean code and best practices.
• Collaborate with cross-functional teams to deliver features from concept to production.
• Debug, profile, and optimize code to improve performance and reliability.
• Use tools and technologies related to: {skills_line}.

Previous Company | Software Developer / Intern | Location | MM/YYYY – MM/YYYY
• Worked on modules or features that contributed directly to business outcomes.
• Wrote maintainable, testable code and participated in code reviews.
• Integrated APIs, databases, or cloud services as required by the project.
"""
    elif category == "data":
        summary = f"""
SUMMARY
{jd_summary_line or f"{job_title} with a strong focus on turning data into actionable insights."}
Hands-on experience in data cleaning, analysis, visualization, and building predictive models.
Comfortable working with large datasets and communicating findings to stakeholders.
"""
        experience_block = f"""
PROFESSIONAL EXPERIENCE
Company Name | {job_title} | Location | MM/YYYY – Present
• Collect, clean, and prepare datasets for analysis and modeling.
• Build dashboards/reports that track key business or product metrics.
• Apply statistical and machine learning techniques to solve business problems.
• Use tools and technologies related to: {skills_line}.

Previous Company | Data Analyst / Intern | Location | MM/YYYY – MM/YYYY
• Assisted in data exploration and visualization for regular reporting.
• Helped stakeholders interpret data and supported decision making.
"""
    elif category == "sales":
        summary = f"""
SUMMARY
{jd_summary_line or f"Results-driven {job_title} with a strong track record in lead generation and deal closure."}
Proven ability to build relationships, understand customer needs, and exceed revenue targets.
Skilled in managing pipelines, handling objections, and closing deals.
"""
        experience_block = f"""
PROFESSIONAL EXPERIENCE
Company Name | {job_title} | Location | MM/YYYY – Present
• Own and manage a sales pipeline from prospecting to closing.
• Conduct product demos, presentations, and negotiations with prospects.
• Consistently achieve or exceed monthly/quarterly sales targets.
• Maintain accurate records in CRM and follow up with clients proactively.

Previous Company | Sales Executive / Inside Sales | Location | MM/YYYY – MM/YYYY
• Generated leads via cold calling, email outreach, and social channels.
• Qualified prospects based on fit, budget, authority, and timeline.
• Supported senior sales staff with proposals and follow-ups.
"""
    elif category == "pm_ba":
        summary = f"""
SUMMARY
{jd_summary_line or f"{job_title} with experience in requirements gathering, stakeholder communication, and delivery."}
Skilled in translating business needs into clear user stories and collaborating with cross-functional teams.
Comfortable managing scope, priorities, and timelines in dynamic environments.
"""
        experience_block = f"""
PROFESSIONAL EXPERIENCE
Company Name | {job_title} | Location | MM/YYYY – Present
• Gather and document business requirements and user needs.
• Define user stories, acceptance criteria, and maintain product backlog.
• Collaborate with engineering, design, and stakeholders to deliver features.
• Track progress, risks, and communicate status transparently.

Previous Company | Business Analyst / Project Coordinator | Location | MM/YYYY – MM/YYYY
• Analyzed processes and identified gaps and opportunities for improvement.
• Supported project planning, tracking, and reporting activities.
"""
    elif category == "design":
        summary = f"""
SUMMARY
{jd_summary_line or f"Creative {job_title} focused on crafting intuitive and visually appealing user experiences."}
Experienced in user research, wireframing, prototyping, and design handoff to engineering teams.
Comfortable iterating based on feedback and usability testing.
"""
        experience_block = f"""
PROFESSIONAL EXPERIENCE
Company Name | {job_title} | Location | MM/YYYY – Present
• Design user interfaces for web/mobile in collaboration with product and engineering.
• Conduct or review user research and usability tests to validate design decisions.
• Create wireframes, prototypes, and design specs using tools like Figma/Sketch.
• Maintain and contribute to design systems and component libraries.

Previous Company | UI/UX Designer / Intern | Location | MM/YYYY – MM/YYYY
• Assisted in designing features and flows for digital products.
• Created visual assets and helped maintain consistent branding.
"""
    elif category == "support":
        summary = f"""
SUMMARY
{jd_summary_line or f"{job_title} focused on delivering excellent customer experiences and efficient issue resolution."}
Experienced in handling tickets, calls, and chats while maintaining high satisfaction scores.
Strong communication, patience, and problem-solving skills.
"""
        experience_block = f"""
PROFESSIONAL EXPERIENCE
Company Name | {job_title} | Location | MM/YYYY – Present
• Respond to customer queries via phone, email, or chat within defined SLAs.
• Troubleshoot issues, coordinate with internal teams, and ensure resolution.
• Maintain detailed case notes and contribute to knowledge base articles.
• Track and report recurring issues or feedback patterns.

Previous Company | Customer Support / Service Desk | Location | MM/YYYY – MM/YYYY
• Handled first-level support, escalating complex issues as needed.
• Assisted in onboarding new users and explaining product features.
"""
    elif category == "finance":
        summary = f"""
SUMMARY
{jd_summary_line or f"{job_title} with experience in financial analysis, reporting, and forecasting."}
Strong analytical skills, attention to detail, and ability to present insights clearly.
Familiar with budgeting, variance analysis, and management reports.
"""
        experience_block = f"""
PROFESSIONAL EXPERIENCE
Company Name | {job_title} | Location | MM/YYYY – Present
• Analyze financial statements, KPIs, and trends to support decision making.
• Assist in preparing budgets, forecasts, and monthly/quarterly reports.
• Build and maintain financial models in Excel / BI tools.
• Work closely with business teams to track spend and performance.

Previous Company | Financial Analyst / Intern | Location | MM/YYYY – MM/YYYY
• Supported financial planning and analysis activities.
• Prepared basic reports and reconciliations under supervision.
"""
    elif category == "hr":
        summary = f"""
SUMMARY
{jd_summary_line or f"{job_title} experienced in recruitment, onboarding, and employee engagement."}
Strong interpersonal skills and understanding of HR processes and policies.
Comfortable partnering with leadership and employees to support people initiatives.
"""
        experience_block = f"""
PROFESSIONAL EXPERIENCE
Company Name | {job_title} | Location | MM/YYYY – Present
• Manage end-to-end recruitment for assigned roles (JD, sourcing, screening, offers).
• Coordinate onboarding, induction, and documentation for new hires.
• Support performance management, feedback cycles, and HR operations.

Previous Company | HR Executive / Recruiter | Location | MM/YYYY – MM/YYYY
• Assisted in scheduling interviews, background checks, and HR documentation.
• Helped organize employee engagement activities and events.
"""
    else:  # generic
        summary = f"""
SUMMARY
{jd_summary_line or f"{job_title} with a strong focus on delivering measurable outcomes and supporting business goals."}
Skilled in collaborating with cross-functional teams, learning quickly, and adapting to new tools and domains.
"""
        experience_block = f"""
PROFESSIONAL EXPERIENCE
Company Name | {job_title} | Location | MM/YYYY – Present
• Describe your main responsibilities and how they relate to {job_title}.
• Highlight 2–4 achievements with measurable impact (revenue, efficiency, satisfaction).
• Mention important tools, systems, or methods you use.

Previous Company | Previous Role | Location | MM/YYYY – MM/YYYY
• Add relevant experience that supports your transition or growth in {job_title}.
"""

    # ==== FINAL TEMPLATE ASSEMBLY ====
    template = f"""
{name_display}
{job_title}
City, Country • Phone • Email • LinkedIn / Portfolio

{summary.strip()}

KEY SKILLS
• {skills_line}

{experience_block.strip()}

PROJECTS
Project Name | Tech/Tools used
• Short description of the project objective and your role.
• Mention specific responsibilities and impact (e.g., metrics improved).

Project Name | Academic / Personal Project
• Describe the problem you solved or value you created.
• Add responsibilities and impact relevant to {job_title}.

EDUCATION
Degree Name (e.g., B.Tech in CSE / BBA / MBA / etc.)
College / University Name | Location | Graduation Year
• Include CGPA / Percentage (if strong and relevant).
• Add coursework relevant to {job_title}.

CERTIFICATIONS & TRAINING
• Certification or Course Name – Platform / Institution – Year
• Short workshops or online courses relevant to this role.

ACHIEVEMENTS
• Awards, recognitions, or performance-based achievements.
• Competitions, hackathons, sales awards, or other highlights.

EXTRACURRICULAR / LEADERSHIP (Optional)
• Leadership roles, volunteering, or organizing activities.

REFERENCES
Available on request.
"""
    return template.strip()


def build_docx_from_template_text(template_text: str) -> bytes:
    """
    Take the generated template text and create a .docx file,
    line by line so the user can easily edit it.
    """
    doc = Document()
    for line in template_text.splitlines():
        doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


# ========= STREAMLIT PAGE ========= #

st.set_page_config(page_title="Resume Analyzer", page_icon="📝", layout="wide")

st.title("📝 Resume Analyzer")

st.markdown(
    "Upload a resume, select a job role, and we'll extract details, show skills inside the resume, "
    "score it against the role, and generate a tailored, role-specific resume template "
    "that you can download as an editable .docx file."
)

# Layout: left controls, right outputs
col_left, col_right = st.columns([1, 2])

with col_left:
    st.subheader("1️⃣ Job & Resume")

    job_title = st.selectbox("Select Job Role", options=list(JOB_PROFILES.keys()))

    candidate_name_input = st.text_input("Candidate Name (optional)", "")

    # 🔹 New: Optional custom JD input
    custom_jd_input = st.text_area(
        "Job Description for this role",
        "",
        help="Paste the exact JD here if you want the analysis & template to use it instead of the default one."
    )

    st.markdown("**Upload Resume (PDF / DOCX / TXT):**")
    uploaded_file = st.file_uploader("", type=["pdf", "docx", "txt"])

    analyze_button = st.button("🔍 Analyze Resume")

with col_right:
    st.subheader("2️⃣ Analysis Output")
    st.info(
        "After you upload a resume and click **Analyze**, we'll show extracted details, "
        "skills found, match score, and a tailored resume template. You can download it as a .docx."
    )

st.markdown("---")

# ========== ANALYSIS LOGIC ========== #

if analyze_button:
    if uploaded_file is None:
        st.error("Please upload a resume file first.")
    else:
        with st.spinner("Analyzing resume..."):
            file_bytes = uploaded_file.read()
            resume_text, image_bytes = extract_text_and_image(file_bytes, uploaded_file.name)

            if not resume_text.strip():
                st.error("Could not read any text from the file. Try another resume or format.")
            else:
                # --- Extract details ---
                details = extract_basic_details(resume_text, candidate_name_input)
                extracted_name = details["name"] or candidate_name_input

                # --- Skills in resume (from global skills universe) ---
                resume_all_skills = extract_skills_from_text(resume_text, ALL_SKILLS)

                # --- JD-specific scoring ---
                jd_profile = JOB_PROFILES[job_title]
                jd_skills = jd_profile["skills"]

                # 🔹 Use custom JD if provided, else default JD
                if custom_jd_input.strip():
                    jd_text = custom_jd_input.strip()
                else:
                    jd_text = jd_profile["jd"]

                resume_skills_for_role = extract_skills_from_text(resume_text, jd_skills)

                matched_skills = resume_skills_for_role
                missing_skills = [s for s in jd_skills if s not in matched_skills]

                total = len(jd_skills) if jd_skills else 1
                score = round(100 * len(matched_skills) / total, 2)

                # ===== SHOW RESULTS =====
                st.subheader("3️⃣ Extracted Candidate Details")

                col_d1, col_d2 = st.columns([1, 2])

                with col_d1:
                    if image_bytes:
                        st.image(image_bytes, caption="Profile photo from resume", width=150)
                    else:
                        st.caption("No profile photo detected (DOCX image only).")

                with col_d2:
                    st.write(f"**Name:** {extracted_name or '—'}")
                    st.write(f"**Email:** {details['email'] or '—'}")
                    st.write(f"**Phone:** {details['phone'] or '—'}")
                    if details["links"]:
                        st.write("**Links:**")
                        for link in details["links"]:
                            st.write(f"- {link}")
                    else:
                        st.write("**Links:** —")

                st.markdown("---")

                st.subheader("4️⃣ Skills Detected in This Resume")
                if resume_all_skills:
                    st.write(", ".join(sorted(set(resume_all_skills))))
                else:
                    st.write("_No known skills from our internal list were detected in this resume._")

                st.markdown("---")
                st.subheader("5️⃣ Role Match Result")

                col_score, col_missing = st.columns([1, 2])

                with col_score:
                    st.markdown("**Match Score for Selected Role:**")
                    st.markdown(f"<h2 style='color:#1f77b4'>{score}%</h2>", unsafe_allow_html=True)

                with col_missing:
                    st.markdown("**Skills you should highlight / add in the resume for this role:**")
                    if missing_skills:
                        st.write(", ".join(missing_skills))
                    else:
                        st.write("_All key skills for this role are already present in the resume._")

                st.markdown("---")
                st.subheader("6️⃣ Tailored Resume Template (Role-Specific)")

                template_text = generate_resume_template(
                    job_title=job_title,
                    candidate_name=extracted_name,
                    matched_skills=matched_skills,
                    missing_skills=missing_skills,
                    jd_text=jd_text,
                )

                # Show template as text (for quick copy/paste)
                st.code(template_text, language="markdown")

                # Build .docx bytes
                docx_bytes = build_docx_from_template_text(template_text)
                base_name = (extracted_name or job_title).replace(" ", "_") or "resume"
                docx_filename = f"{base_name}_resume_template.docx"

                # Download button for .docx
                st.download_button(
                    label="⬇️ Download Resume Template (.docx)",
                    data=docx_bytes,
                    file_name=docx_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

                # ===== Save to results.json for dashboard =====
                os.makedirs("data", exist_ok=True)
                results_file = os.path.join("data", "results.json")

                record = {
                    "timestamp": datetime.now().isoformat(),
                    "candidate_name": extracted_name or uploaded_file.name,
                    "job_title": job_title,
                    "score": score,
                    "jd_skills": jd_skills,
                    "resume_skills": matched_skills,
                    "missing_skills": missing_skills,
                }

                append_result_to_json(results_file, record)

                st.success("Analysis saved to dashboard data. Open the main Dashboard to view stats and trends.")

footer_html = """
<style>
footer {
    visibility: hidden;
}
#custom-footer {
    position: scroll;
    left: 0;
    bottom: 0;
    width: 100%;
    background: linear-gradient(90deg,#1f77b4,#b8860b);
    color: white;
    text-align: center;
    padding: 8px 0;
    font-size: 14px;
    font-weight: 500;
    letter-spacing: 0.5px;
    z-index: 9999;
}
</style>

<div id="custom-footer">
    Developed by Akhil • ✉️akhillade431@gmail.com
</div>
"""
st.markdown(footer_html, unsafe_allow_html=True)
